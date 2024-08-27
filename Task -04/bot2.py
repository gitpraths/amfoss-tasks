import csv
import io
from typing import List
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, CallbackContext, CallbackQueryHandler, ConversationHandler
from docx import Document

# Your bot token and Google API key
BOT_TOKEN = '6790370513:AAElipW4HYV6RLEq7bDEpF06dF1z0wfXv2k'
GOOGLE_API_KEY = 'AIzaSyDnk0Y-6FEfGLYH7NqrBaue-7vAPS6HqzU'

# Globals
reading_list = []

# States for conversation
GENRE, BOOK_NAME, BOOK_DETAILS, BOOK_TO_DELETE = range(4)

async def start(update: Update, context: CallbackContext):
    await update.message.reply_text('Welcome! Use /book to get book details by genre, /preview to get a preview link, /list to manage your reading list, or /help for more information.')

async def help_command(update: Update, context: CallbackContext):
    help_text = (
        "/start - Start the bot\n"
        "/book - Get book details by genre\n"
        "/preview - Get a preview link for a book\n"
        "/list - Manage your reading list\n"
        "/help - Get help information"
    )
    await update.message.reply_text(help_text)

async def book(update: Update, context: CallbackContext):
    await update.message.reply_text('Please enter the genre:')
    return GENRE

async def handle_genre(update: Update, context: CallbackContext):
    genre = update.message.text
    books = fetch_books_by_genre(genre)
    
    # Create CSV file
    csv_file = io.StringIO()
    csv_writer = csv.writer(csv_file)
    csv_writer.writerow(['Title', 'Author', 'Description', 'Preview Link'])
    
    for book in books:
        csv_writer.writerow([book['title'], book['authors'], book['description'], book['preview_link']])
    
    csv_file.seek(0)
    await update.message.reply_document(document=csv_file, filename='books.csv')
    return ConversationHandler.END

async def preview(update: Update, context: CallbackContext):
    await update.message.reply_text('Please enter the book name:')
    return BOOK_NAME

async def handle_book_name(update: Update, context: CallbackContext):
    book_name = update.message.text
    preview_link = fetch_preview_link(book_name)
    
    if preview_link:
        await update.message.reply_text(f'Preview link: {preview_link}')
    else:
        await update.message.reply_text('No preview link found for this book.')
    return ConversationHandler.END

async def list_command(update: Update, context: CallbackContext):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add')],
        [InlineKeyboardButton("Delete a book", callback_data='delete')],
        [InlineKeyboardButton("View Reading List", callback_data='view')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text('Choose an option:', reply_markup=reply_markup)

async def handle_button(update: Update, context: CallbackContext):
    query = update.callback_query
    choice = query.data
    
    await query.answer()

    if choice == 'add':
        await query.edit_message_text(text="Send me the book details to add.")
        return BOOK_DETAILS
    
    elif choice == 'delete':
        await query.edit_message_text(text="Send me the book title to delete.")
        return BOOK_TO_DELETE
    
    elif choice == 'view':
        docx_file = generate_reading_list_docx(reading_list)
        await query.message.reply_document(document=docx_file, filename='reading_list.docx')
        return ConversationHandler.END

async def handle_book_details(update: Update, context: CallbackContext):
    book_details = update.message.text
    reading_list.append(book_details)
    await update.message.reply_text("Book added to the reading list.")
    return ConversationHandler.END

async def handle_book_to_delete(update: Update, context: CallbackContext):
    book_title = update.message.text
    global reading_list
    reading_list = [book for book in reading_list if book_title not in book]
    await update.message.reply_text("Book removed from the reading list.")
    return ConversationHandler.END

def fetch_books_by_genre(genre: str) -> List[dict]:
    service = build('books', 'v1', developerKey=GOOGLE_API_KEY)
    try:
        response = service.volumes().list(q=f'subject:{genre}', maxResults=10).execute()
        items = response.get('items', [])
        books = [{
            'title': item['volumeInfo'].get('title', 'No Title'),
            'authors': ', '.join(item['volumeInfo'].get('authors', ['Unknown Author'])),
            'description': item['volumeInfo'].get('description', 'No Description'),
            'preview_link': item['volumeInfo'].get('previewLink', 'No Preview Link')
        } for item in items]
        return books
    except HttpError as error:
        print(f"An error occurred: {error}")
        return []

def fetch_preview_link(book_name: str) -> str:
    service = build('books', 'v1', developerKey=GOOGLE_API_KEY)
    try:
        response = service.volumes().list(q=book_name, maxResults=1).execute()
        items = response.get('items', [])
        if items:
            preview_link = items[0]['volumeInfo'].get('previewLink', 'No Preview Link')
            return preview_link
        return 'No Preview Link'
    except HttpError as error:
        print(f"An error occurred: {error}")
        return 'No Preview Link'

def generate_reading_list_docx(reading_list: List[str]) -> io.BytesIO:
    doc = Document()
    doc.add_heading('Reading List', level=1)
    for entry in reading_list:
        doc.add_paragraph(entry)
    doc_file = io.BytesIO()
    doc.save(doc_file)
    doc_file.seek(0)
    return doc_file

def main():
    application = Application.builder().token(BOT_TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[
            CommandHandler('book', book),
            CommandHandler('preview', preview),
            CommandHandler('list', list_command),
        ],
        states={
            GENRE: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_genre)],
            BOOK_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_book_name)],
            BOOK_DETAILS: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_book_details)],
            BOOK_TO_DELETE: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_book_to_delete)],
        },
        fallbacks=[CommandHandler('start', start)],
    )

    application.add_handler(CommandHandler('start', start))
    application.add_handler(CommandHandler('help', help_command))
    application.add_handler(conv_handler)
    application.add_handler(CallbackQueryHandler(handle_button))

    application.run_polling()

if __name__ == '__main__':
    main()